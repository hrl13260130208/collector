#!/usr/bin/python
#-*- coding: utf-8 -*-

from  docx import Document



def create_arff(file_path,attrs,name):
    arff_file=open("C:/temp/"+name+".arff","w+")
    arff_file.write("@relation order\n")
    for attr in attrs:
        arff_file.write("@attribute "+attr+" {'0','1'}\n")
    arff_file.write("@data\n")

    for line in open(file_path).readlines():
        # print(line)
        write_line=""
        keys = line.replace("\n", "").split(",")
        for attr in attrs:
            find=False
            for key in keys:
                if key == "":
                    continue
                if attr==key:
                    write_line+="'1',"
                    find=True
            if not find:
                write_line += "'0',"
        arff_file.write(write_line[:-1]+"\n")
        # print(write_line[:-1])

if __name__ == '__main__':
    path = "C:/temp/part-r-2017"
    attrs = {}
    for line in open(path).readlines():
        keys = line.replace("\n", "").split(",")
        for key in keys:
            if key == "":
                continue
            if key in attrs:
                attrs[key] += 1
            else:
                attrs[key] = 1
    # print(attrs)
    sort_attrs = sorted(attrs.items(), key=lambda d: d[1], reverse=True)
    # print(sort_attrs)
    attr_list = []
    for attr in sort_attrs:
        print(type(attr), attr, attr[0])
        attr_list.append(attr[0])

    create_arff(path,attr_list,"test")



# doc=Document("C:/temp/pdf_test/00GcBkLKczcf.docx")
# for p in doc.paragraphs:
#     print(p.text)

# file=open("C:/temp/新建文本文档.txt")
# strings=""
# for line in file.readlines():
#     strings+="\""+line.replace("\n","")+"\","
# print(strings)
# #打开一个pdf文件
# fp = open("C:/File/GOtNEOBCKt5N.pdf", 'rb')
# #创建一个PDF文档解析器对象
# parser = PDFParser(fp)
# #创建一个PDF文档对象存储文档结构
# #提供密码初始化，没有就不用传该参数
# #document = PDFDocument(parser, password)
# document = PDFDocument(parser)
# #检查文件是否允许文本提取
# if not document.is_extractable:
#     raise PDFTextExtractionNotAllowed
# #创建一个PDF资源管理器对象来存储共享资源
# #caching = False不缓存
# rsrcmgr = PDFResourceManager(caching = False)
# # 创建一个PDF设备对象
# laparams = LAParams()
# # 创建一个PDF页面聚合对象
# device = PDFPageAggregator(rsrcmgr, laparams=laparams)
# #创建一个PDF解析器对象
# interpreter = PDFPageInterpreter(rsrcmgr, device)
# #处理文档当中的每个页面
#
# # doc.get_pages() 获取page列表
# #for i, page in enumerate(document.get_pages()):
# #PDFPage.create_pages(document) 获取page列表的另一种方式
# replace=re.compile(r'\s+')
# # 循环遍历列表，每次处理一个page的内容
# for page in PDFPage.create_pages(document):
#     interpreter.process_page(page)
#     # 接受该页面的LTPage对象
#     layout=device.get_result()
#     # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
#     # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
#     for x in layout:
#         #如果x是水平文本对象的话
#         if(isinstance(x,LTTextBoxHorizontal)):
#             text=re.sub(replace,'',x.get_text())
#             if len(text)>100:
#                 print (text.__len__(),text)